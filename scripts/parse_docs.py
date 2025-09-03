from docx import Document

# Load the document
document = Document('your_document.docx')

# Extract paragraphs
for para in document.paragraphs:
    print(para.text)

# Extract text from tables
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                print(para.text)

# Extract images (metadata like width/height)
for shape in document.inline_shapes:
    if shape.type == 3:  # 3 corresponds to INLINE_PICTURE
        print(f"Image dimensions: {shape.width.emu / 914400} inches x {shape.height.emu / 914400} inches")